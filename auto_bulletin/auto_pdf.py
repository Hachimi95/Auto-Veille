"""
Bulletin DOCX/PDF generator.

Key fixes vs previous version:
  - Placeholder replacement now merges split runs before matching,
    so [CVE2], [score], [risques], [Delai], [Ref], [Mitigations] are
    found even when Word splits them across multiple <w:r> elements.
  - All paragraphs are visited recursively: body paragraphs, every
    table cell (including inside text-boxes / mc:Choice blocks), and
    header/footer paragraphs.
  - fix_table_properties targets the correct body table by scanning
    for [CVE] rather than assuming a fixed index.
  - import re moved to module level (was inside a loop).
  - The version-splitting workaround that duplicated logic is removed;
    split_version_text handles everything.
"""

import os
import re
import json
import platform
import shutil
import subprocess
import tempfile
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

__all__ = ["generate_pdf_from_json", "generate_docx_from_json"]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

MONTHS_FR = {
    "janvier": "01", "février": "02", "mars": "03", "avril": "04",
    "mai": "05", "juin": "06", "juillet": "07", "août": "08",
    "septembre": "09", "octobre": "10", "novembre": "11", "décembre": "12",
}

VERSION_PATTERN = re.compile(
    r'\d+\.\d+\.\d+\.\d+\/\.\d+'
    r'|\d+\.\d+\.\d+\+security-\d{2}rc\d+'
    r'|\d+\.\d+\.\d+\+security-\d{2}'
    r'|\d+\.\d+\.\d+\.\d+'
    r'|\d+\.\d+\.\d+rc\d+'
    r'|\d+\.\d+\.\d+'
    r'|\d{1,2}\.\d{1,2}\.x'
    r'|\d+\.x'
    r'|v\d+\.\d+'
    r'|\d{1,2}\.\d{1,2}'
)


def convert_date_format(french_date: str) -> str:
    """Convert French date string to dd/mm/yyyy."""
    try:
        parts = french_date.split()
        day, month_fr, year = parts[0], parts[1].lower(), parts[2]
        month = MONTHS_FR.get(month_fr, "")
        return f"{day}/{month}/{year}" if month else french_date
    except Exception:
        return french_date


def split_version_text(text: str):
    """Return list of (fragment, is_bold) where version numbers are bold."""
    parts, last = [], 0
    for m in VERSION_PATTERN.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((text[m.start():m.end()], True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts


def _copy_run_format(src_run, dst_run):
    """Copy font attributes from src_run to dst_run (best-effort)."""
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    dst_run.font.bold = src_run.font.bold
    try:
        if src_run.font.color and src_run.font.color.type is not None:
            dst_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Paragraph-level placeholder replacement
# ---------------------------------------------------------------------------

def _merge_paragraph_runs(paragraph) -> str:
    """Return the full text of a paragraph by joining all run texts."""
    return "".join(r.text or "" for r in paragraph.runs)


def replace_placeholders_in_paragraph(paragraph, placeholders: dict):
    """
    Replace any placeholder found in *paragraph*.

    Word often splits a single placeholder like [CVE2] across several
    consecutive runs, e.g. ['[', 'CVE', '2]'].  We detect this by
    reading the concatenated run text.  When a match is found we clear
    all runs and rewrite the paragraph content.
    """
    full_text = _merge_paragraph_runs(paragraph)
    if not full_text:
        return

    matched_key = next((k for k in placeholders if k in full_text), None)
    if matched_key is None:
        return

    # Preserve formatting from the first non-empty run before clearing.
    original_runs = list(paragraph.runs)
    ref_run = next((r for r in original_runs if r.text.strip()), None)

    value = placeholders[matched_key]

    # --- [CVE] / [CVE2] ---
    if matched_key in ("[CVE]", "[CVE2]"):
        paragraph.clear()
        cves = str(value).split("\n")
        n = len(cves)
        if   n <= 6:  font_size, line_spacing, space_before = Pt(16), 1.6, Pt(70)
        elif n <= 10: font_size, line_spacing, space_before = Pt(15), 1.5, Pt(50)
        elif n <= 15: font_size, line_spacing, space_before = Pt(14), 1.4, Pt(40)
        elif n <= 20: font_size, line_spacing, space_before = Pt(13), 1.3, Pt(30)
        elif n <= 25: font_size, line_spacing, space_before = Pt(12), 1.1, Pt(20)
        elif n <= 30: font_size, line_spacing, space_before = Pt(11), 1.0, Pt(20)
        elif n <= 35: font_size, line_spacing, space_before = Pt(11), 0.8, Pt(10)
        elif n <= 40: font_size, line_spacing, space_before = Pt(11), 0.8, Pt(5)
        elif n <= 45: font_size, line_spacing, space_before = Pt(10), 0.5, Pt(0)
        else:         font_size, line_spacing, space_before = Pt(9),  0.1, Pt(0)

        for i, cve in enumerate(cves):
            run = paragraph.add_run(cve.strip())
            if ref_run:
                _copy_run_format(ref_run, run)
            run.font.size = font_size
            if i < n - 1:
                paragraph.add_run("\n")

        paragraph.paragraph_format.space_after  = Pt(4)
        paragraph.paragraph_format.space_before = space_before
        paragraph.paragraph_format.line_spacing = line_spacing

    # --- [Produits affectés] ---
    elif matched_key == "[Produits affectés]":
        align = paragraph.alignment
        style = paragraph.style
        paragraph.clear()

        products = value if isinstance(value, list) else [str(value)]
        for i, product in enumerate(products):
            bullet = paragraph.add_run(chr(183) + "   ")
            bullet.font.name = "Symbol"
            bullet.font.size = Pt(11)
            for fragment, bold in split_version_text(product):
                r = paragraph.add_run(fragment)
                r.font.name = "Arial"
                r.font.size = Pt(10)
                r.font.bold = bold
            if i < len(products) - 1:
                paragraph.add_run("\n")

        paragraph.paragraph_format.line_spacing = 1.6
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after  = Pt(1)
        paragraph.alignment = align
        paragraph.style     = style

    # --- [Mitigations] ---
    elif matched_key == "[Mitigations]":
        align = paragraph.alignment
        style = paragraph.style
        paragraph.clear()

        mitigations = value if isinstance(value, list) else [value]
        for i, mitigation in enumerate(mitigations):
            details = _parse_mitigation(mitigation)
            rec_text = (details.get("recommendation") or "").strip()
            versions = _expand_versions(details.get("versions") or [])

            if rec_text:
                r = paragraph.add_run(rec_text)
                r.font.name = "Arial"
                r.font.size = Pt(10)
                r.font.bold = False
                if versions:
                    paragraph.add_run("\n")

            for version in versions:
                v = str(version).strip()
                if not v:
                    continue
                indent = paragraph.add_run("          ")
                indent.font.name = "Arial"
                indent.font.size = Pt(10)
                bullet = paragraph.add_run(chr(183) + "   ")
                bullet.font.name = "Symbol"
                bullet.font.size = Pt(11)
                for fragment, bold in split_version_text(v):
                    r = paragraph.add_run(fragment)
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
                    r.font.bold = bold
                paragraph.add_run("\n")

            if i < len(mitigations) - 1:
                paragraph.add_run("\n")

        paragraph.paragraph_format.line_spacing = 1.6
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after  = Pt(1)
        paragraph.alignment = align
        paragraph.style     = style

    # --- all other placeholders (simple text replacement) ---
    else:
        new_text = full_text.replace(matched_key, str(value))
        paragraph.clear()
        run = paragraph.add_run(new_text)
        if ref_run:
            _copy_run_format(ref_run, run)


def _parse_mitigation(mitigation) -> dict:
    """Normalise a mitigation entry into {'recommendation': str, 'versions': list}."""
    if isinstance(mitigation, str):
        try:
            parsed = json.loads(mitigation)
            if isinstance(parsed, dict):
                if "recommendation" in parsed:
                    return parsed
                # single-key wrapper: {product: {recommendation, versions}}
                return next(iter(parsed.values()), {})
        except json.JSONDecodeError:
            pass
        return {"recommendation": mitigation, "versions": []}

    if isinstance(mitigation, dict):
        if "recommendation" in mitigation:
            return mitigation
        return next(iter(mitigation.values()), mitigation)

    return {"recommendation": str(mitigation), "versions": []}


def _expand_versions(versions: list) -> list:
    """Split version strings that contain multiple concatenated versions."""
    result = []
    for v in versions:
        v = str(v).strip()
        if not v:
            continue
        if v.count("ou ultérieure") > 1:
            parts = re.split(r"(ou ultérieure)", v)
            buf = ""
            for part in parts:
                buf += part
                if part == "ou ultérieure":
                    result.append(buf.strip())
                    buf = ""
            if buf.strip():
                result.append(buf.strip())
        elif v.count("or later") > 1:
            parts = re.split(r"(or later)", v)
            buf = ""
            for part in parts:
                buf += part
                if part == "or later":
                    result.append(buf.strip())
                    buf = ""
            if buf.strip():
                result.append(buf.strip())
        else:
            result.append(v)
    return result


# ---------------------------------------------------------------------------
# Document-wide paragraph iteration (body + all tables + text boxes)
# ---------------------------------------------------------------------------

def iter_all_paragraphs(doc):
    """
    Yield every paragraph in the document:
      - body paragraphs
      - cells in every table (including tables inside text boxes / mc:Choice)
      - headers and footers
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Collect all <w:p> elements under <w:body> — this covers body paragraphs,
    # table cells, AND text-box content (txbxContent) in one pass.
    # We skip <mc:Fallback> blocks to avoid double-processing duplicate tables.
    MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    fallbacks = set(
        id(el) for el in doc.element.body.iter(f"{{{MC}}}Fallback")
    )

    from docx.text.paragraph import Paragraph as DocxParagraph

    for p_elem in doc.element.body.iter(f"{{{W}}}p"):
        # Skip if this <w:p> is inside a Fallback block
        if any(id(anc) in fallbacks for anc in p_elem.iterancestors()):
            continue
        yield DocxParagraph(p_elem, doc)

    # Headers and footers
    for section in doc.sections:
        for hdr_ftr in (
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ):
            if hdr_ftr is None:
                continue
            for p_elem in hdr_ftr._element.iter(f"{{{W}}}p"):
                yield DocxParagraph(p_elem, hdr_ftr)


# ---------------------------------------------------------------------------
# Table row-height fix
# ---------------------------------------------------------------------------

def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height_pt))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def fix_table_properties(doc):
    """
    Find the main content table (the one that contains [CVE]) and set
    a fixed height on its second row so content fits properly.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    fallbacks = set(id(el) for el in doc.element.body.iter(f"{{{MC}}}Fallback"))

    from docx.table import Table

    for tbl_elem in doc.element.body.iter(f"{{{W}}}tbl"):
        # Skip fallback copies
        if any(id(anc) in fallbacks for anc in tbl_elem.iterancestors()):
            continue
        texts = [t.text or "" for t in tbl_elem.iter(f"{{{W}}}t")]
        if "[CVE]" in texts or any("[CVE]" in t for t in texts):
            tbl = Table(tbl_elem, doc)
            if len(tbl.rows) >= 2:
                set_row_height(tbl.rows[1], 7800)
            break  # only fix the first matching table


# ---------------------------------------------------------------------------
# LibreOffice helpers
# ---------------------------------------------------------------------------

def _find_libreoffice() -> str | None:
    env = os.getenv("SOFFICE_PATH")
    candidates = [env] if env else []
    for name in ("soffice", "libreoffice", "lowriter"):
        p = shutil.which(name)
        if p:
            candidates.append(p)
    candidates += [
        "/usr/bin/soffice", "/usr/bin/libreoffice",
        "/usr/local/bin/soffice", "/usr/local/bin/libreoffice",
        "/usr/lib/libreoffice/program/soffice",
        "/snap/bin/libreoffice", "/opt/libreoffice/program/soffice",
    ]
    seen = set()
    for c in candidates:
        if not c or c in seen:
            continue
        seen.add(c)
        if os.path.isfile(c) and os.access(c, os.X_OK):
            return c
    return None


def convert_docx_to_pdf_libreoffice(docx_path: str) -> str:
    soffice = _find_libreoffice()
    if not soffice:
        raise RuntimeError(
            "LibreOffice not found. Install with:\n"
            "  sudo apt-get install -y libreoffice-core libreoffice-writer fonts-dejavu"
        )
    out_dir = os.path.dirname(docx_path)
    base = os.path.splitext(os.path.basename(docx_path))[0]
    proc = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True,
        check=True, timeout=120,
    )
    expected = os.path.join(out_dir, f"{base}.pdf")
    if os.path.exists(expected):
        return expected
    for f in os.listdir(out_dir):
        if f.lower().endswith(".pdf") and f.startswith(base):
            return os.path.join(out_dir, f)
    raise RuntimeError(f"Conversion succeeded but PDF not found for: {base}")


# ---------------------------------------------------------------------------
# Main public API
# ---------------------------------------------------------------------------

def generate_docx_from_json(json_path: str, bulletin_id: str, out_dir: str = None) -> str:
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        raise ValueError("JSON must be a top-level object.")

    # Build output filename
    date_raw = data.get("Date", "")
    if date_raw:
        parts = date_raw.split()
        if len(parts) >= 3:
            day = parts[0]
            month = MONTHS_FR.get(parts[1].lower(), "00")
            year = parts[2]
            formatted_date = f"{day}{month}{year}"
        else:
            formatted_date = datetime.now().strftime("%d%m%Y")
    else:
        formatted_date = datetime.now().strftime("%d%m%Y")

    titre = data.get("titre", "Unknown_Advisory")
    raw_name = f"{formatted_date}-{bulletin_id} - {titre}"
    base_name = "".join(c for c in raw_name if c.isalnum() or c in "-_ ").rstrip()

    # Load template
    tpl_dir  = os.path.dirname(os.path.abspath(__file__))
    tpl_path = os.path.join(tpl_dir, "template5.docx")
    doc = Document(tpl_path if os.path.exists(tpl_path) else None)

    # Build placeholder map
    display_date = convert_date_format(date_raw) if date_raw else ""
    placeholders = {
        "[Titre]":            data.get("titre", ""),
        "[titre]":            data.get("titre", ""),
        "[CVE2]":             "\n".join(data.get("CVEs ID", [])),
        "[CVE]":              "\n".join(data.get("CVEs ID", [])),
        "[Produits affectés]": data.get("Produits affectés", []),
        "[Description]":      data.get("Description", ""),
        "[Exploit]":          data.get("Exploit", ""),
        "[Delai]":            data.get("Delai", ""),
        "[score]":            data.get("score", ""),
        "[Date]":             display_date,
        "[Ref]":              "\n".join(data.get("Références", [])),
        "[Mitigations]":      data.get("Mitigations", []),
        "[risques]":          "\n".join(
            r + "\n-" for r in data.get("risques", [])
        ).rstrip("\n-"),
    }

    # Fix row heights before replacing content
    fix_table_properties(doc)

    # Replace every paragraph everywhere in the document
    for paragraph in iter_all_paragraphs(doc):
        replace_placeholders_in_paragraph(paragraph, placeholders)

    # Save
    if out_dir is None:
        out_dir = tpl_dir
    os.makedirs(out_dir, exist_ok=True)
    docx_path = os.path.join(out_dir, f"{base_name}.docx")
    doc.save(docx_path)
    return docx_path


def generate_pdf_from_json(json_path: str, bulletin_id: str, return_bytes: bool = False):
    """Generate PDF (and optionally DOCX bytes) from a JSON advisory file."""
    tmpdir = tempfile.mkdtemp(prefix="bulletin_")
    try:
        docx_path = generate_docx_from_json(json_path, bulletin_id, out_dir=tmpdir)

        pdf_path = None
        try:
            if platform.system() == "Windows":
                try:
                    from win32com import client as win32
                    word = win32.Dispatch("Word.Application")
                    word.Visible = False
                    wb = word.Documents.Open(os.path.abspath(docx_path))
                    pdf_path = docx_path.replace(".docx", ".pdf")
                    wb.SaveAs(pdf_path, FileFormat=17)
                    wb.Close()
                    word.Quit()
                except Exception:
                    pdf_path = convert_docx_to_pdf_libreoffice(docx_path)
            else:
                pdf_path = convert_docx_to_pdf_libreoffice(docx_path)
        except Exception as e:
            print(f"[bulletin] PDF conversion failed: {e}")

        if not return_bytes:
            # Move files out of tmpdir to the script directory and return pdf path
            dest_dir = os.path.dirname(os.path.abspath(__file__))
            final_docx = shutil.copy2(docx_path, dest_dir)
            if pdf_path and os.path.exists(pdf_path):
                final_pdf = shutil.copy2(pdf_path, dest_dir)
                return final_pdf
            return final_docx

        # return_bytes mode
        result = {}
        if pdf_path and os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                result["pdf_bytes"] = f.read()
            result["pdf_name"] = os.path.basename(pdf_path)
        else:
            result["pdf_bytes"] = None
            result["pdf_name"] = None

        if os.path.exists(docx_path):
            with open(docx_path, "rb") as f:
                result["docx_bytes"] = f.read()
            result["docx_name"] = os.path.basename(docx_path)
        else:
            result["docx_bytes"] = None
            result["docx_name"] = None

        return result

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)